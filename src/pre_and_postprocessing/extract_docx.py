from docx import Document
import re
import os
from model.run_request import get_best_answer

# Keywords for identifying questions
default_keywords = [
    "Erklären", "Beschreiben", "Formulieren", "Liefern", "Diskutieren",
    "Veranschaulichen", "Klären", "Definieren", "Umreißen", "Zusammenfassen",
    "Begründen", "Vergleichen", "Kontrastieren", "Analysieren", "Bewerten",
    "Ausarbeiten", "Identifizieren", "Erkunden", "Interpretieren", "Untersuchen"
]


# Function to get keywords from the user and merge with default keywords
def get_keywords():
    user_keywords = input("Please enter additional keywords, separated by commas: ").split(',')
    user_keywords = [keyword.strip() for keyword in user_keywords]
    return default_keywords + user_keywords


def convert_doc_to_docx(doc_path):
    docx_path = doc_path.replace('.doc', '.docx')
    word = wc.Dispatch('Word.Application')
    doc = word.Documents.Open(doc_path)
    doc.SaveAs(docx_path, 12)  # 12 stands for .docx file format
    doc.Close()
    word.Quit()
    return docx_path


# Function to extract questions from a Word document
def extract_questions_docx(doc_path, keywords):
    avoid_dups = set()
    doc = Document(doc_path)
    questions = []
    counter = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        # print(text)
        if '?' in text:
            print(f"Question: {text}")
            questions.append(text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if '?' in cell.text:
                        counter += 1
                        if cell.text not in avoid_dups:
                            print(cell.text)
                            questions.append(cell.text)
                            avoid_dups.add(cell.text)

    return questions

def extract_questions_from_docx(file_path):
    # Load the DOCX file
    document = Document(file_path)
    
    questions = []

    if not os.path.exists(file_path):
        print(f"Error: The file {file_path} does not exist.")
        return []

    # Iterate through all paragraphs in the document
    for paragraph in document.paragraphs:
        # Get the text of the paragraph and strip leading/trailing whitespace
        text = paragraph.text.strip()
        if text:
            # Use regex to find all questions in the paragraph
            found_questions = re.findall(r'[^.?!]*\?+', text)
            questions.extend(found_questions)
    
    return questions

def docx_answer_questions(file_path):
    questions = extract_questions_from_docx(file_path)

    if not questions:
        print(f"No questions found in file: {file_path}.")
        return
    
    os.makedirs(os.path.dirname("./answered_questions.txt"), exist_ok=True)
    with open("./answered_questions.txt", 'w', encoding='utf-8') as file:
   
        for i, question in enumerate(questions, 1):
            answer, src = get_best_answer(question)
            file.write(f"Question {i}: {question}\n")
            file.write(f"\tAnswer: {answer}, {src}.\n\n")


    
    



#original_doc_path2 = r'C:\Users\jansu\OneDrive\Desktop\DN+ DS\DS Group Project\Organ\Data\Data\Fragebogen - Buchfuhrung und Logistiksystem.doc'
#original_doc_path = r'C:\Users\jansu\OneDrive\Desktop\DN+ DS\DS Group Project\Organ\Data\Data\Fragebogen - Buchfuhrung und Logistiksystem.docx'
# Convert .doc to .docx

#if ".docx" not in original_doc_path:
#    print("test")
#    original_doc_path = convert_doc_to_docx(original_doc_path2)

# Extract questions from the converted .docx file
# questions = extract_questions_docx(original_doc_path, default_keywords)

# Display the extracted questions
# for question in questions:
#    print(question)

# Clean up any intermediate files if desired
# os.remove(docx_path)
